import pdfplumber
import docx
import pytesseract
import pypdfium2 as pdfium
from PIL import Image
from io import BytesIO
from utils.errors import ExtractionError

def validate_file(file_content: bytes, content_type: str):
    if len(file_content) > 10 * 1024 * 1024:
        raise ExtractionError("File size exceeds the 10MB limit.")
        
    allowed_types = [
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]
    if content_type not in allowed_types:
        raise ExtractionError(f"Unsupported file format: {content_type}. Only PDF and DOCX are allowed.")

def extract_pdf(file_bytes: bytes) -> str:
    text = ""
    try:
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        return text
    except Exception as e:
        raise ExtractionError(f"Native PDF extraction failed: {str(e)}")

def is_likely_scanned(extracted_text: str) -> bool:
    return len(extracted_text.strip()) < 100

def ocr_fallback(file_bytes: bytes) -> str:
    try:
        # Load PDF using pypdfium2 (safer and cleaner dependency compared to pdf2image + poppler)
        pdf = pdfium.PdfDocument(file_bytes)
        text = ""
        for i in range(len(pdf)):
            page = pdf.get_page(i)
            # Render page to a bitmap (scale=2 for high resolution for OCR accuracy)
            bitmap = page.render(scale=2)
            pil_image = bitmap.to_pil()
            
            # Execute OCR using pytesseract
            page_text = pytesseract.image_to_string(pil_image)
            if page_text:
                text += page_text + "\n"
                
            page.close()
        pdf.close()
        return text
    except Exception as e:
        raise ExtractionError(f"OCR translation failed: {str(e)}")

def extract_docx(file_bytes: bytes) -> str:
    try:
        doc = docx.Document(BytesIO(file_bytes))
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        raise ExtractionError(f"Word document extraction failed: {str(e)}")

def extract_text(file_bytes: bytes, content_type: str) -> str:
    validate_file(file_bytes, content_type)
    
    if content_type == "application/pdf":
        raw_text = extract_pdf(file_bytes)
        if is_likely_scanned(raw_text):
            return ocr_fallback(file_bytes)
        return raw_text
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_docx(file_bytes)
    else:
        raise ExtractionError("Invalid file type routed to extractor.")
